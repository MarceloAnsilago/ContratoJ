from io import BytesIO
from html import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer


BOLD_START = "[[B]]"
BOLD_END = "[[/B]]"


def _linhas(texto: str) -> list[str]:
    return texto.replace("\r\n", "\n").replace("\r", "\n").split("\n")


def _partes_formatadas(texto: str) -> list[tuple[str, bool]]:
    partes = []
    posicao = 0

    while posicao < len(texto):
        inicio = texto.find(BOLD_START, posicao)
        if inicio == -1:
            partes.append((texto[posicao:], False))
            break

        if inicio > posicao:
            partes.append((texto[posicao:inicio], False))

        conteudo_inicio = inicio + len(BOLD_START)
        fim = texto.find(BOLD_END, conteudo_inicio)
        if fim == -1:
            partes.append((texto[inicio:], False))
            break

        partes.append((texto[conteudo_inicio:fim], True))
        posicao = fim + len(BOLD_END)

    return [(conteudo, negrito) for conteudo, negrito in partes if conteudo]


def _texto_pdf_formatado(texto: str) -> str:
    trechos = []
    for conteudo, negrito in _partes_formatadas(texto):
        conteudo_escapado = escape(conteudo)
        if negrito:
            trechos.append(f"<b>{conteudo_escapado}</b>")
        else:
            trechos.append(conteudo_escapado)
    return "".join(trechos)


def _adicionar_runs_formatados(paragrafo, texto: str, *, negrito_base: bool, tamanho: int) -> None:
    for conteudo, negrito in _partes_formatadas(texto):
        run = paragrafo.add_run(conteudo)
        run.bold = negrito_base or negrito
        run.font.name = "Courier New"
        run.font.size = Pt(tamanho)


def _linha_dupla_pdf() -> list:
    return [
        HRFlowable(width="55%", thickness=0.8, color="#111111", spaceBefore=0, spaceAfter=2, hAlign="CENTER"),
        HRFlowable(width="55%", thickness=0.8, color="#111111", spaceBefore=0, spaceAfter=10, hAlign="CENTER"),
    ]


def _adicionar_linha_dupla_docx(documento: Document) -> None:
    for indice in range(2):
        paragrafo = documento.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo.paragraph_format.space_after = Pt(0 if indice == 0 else 10)
        run = paragrafo.add_run("________________________________________")
        run.font.name = "Courier New"
        run.font.size = Pt(8)


def gerar_contrato_pdf(texto: str, titulo: str = "Contrato") -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=titulo,
    )

    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "ContratoNormal",
        parent=styles["Normal"],
        fontName="Courier",
        fontSize=12,
        leading=18,
        alignment=4,
        firstLineIndent=1.25 * cm,
        spaceAfter=10,
    )
    heading = ParagraphStyle(
        "ContratoTitulo",
        parent=styles["Title"],
        fontName="Courier-Bold",
        fontSize=14,
        leading=18,
        alignment=1,
        spaceAfter=5,
    )

    elementos = []
    primeira_linha = True
    for linha in _linhas(texto):
        bloco = linha.strip()
        if not bloco:
            elementos.append(Spacer(1, 8))
            continue

        indice = 0 if primeira_linha else 1
        estilo = heading if indice == 0 else normal
        elementos.append(Paragraph(_texto_pdf_formatado(bloco), estilo))
        if indice == 0:
            elementos.extend(_linha_dupla_pdf())
            primeira_linha = False

    doc.build(elementos)
    buffer.seek(0)
    return buffer.getvalue()


def gerar_contrato_docx(texto: str, titulo: str = "Contrato") -> bytes:
    documento = Document()
    documento.core_properties.title = titulo

    estilo = documento.styles["Normal"]
    estilo.font.name = "Courier New"
    estilo.font.size = Pt(12)

    primeira_linha = True
    for linha in _linhas(texto):
        bloco = linha.strip()
        paragrafo = documento.add_paragraph()
        if not bloco:
            paragrafo.paragraph_format.space_after = Pt(4)
            continue

        indice = 0 if primeira_linha else 1
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER if indice == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafo.paragraph_format.first_line_indent = Pt(0 if indice == 0 else 36)
        paragrafo.paragraph_format.space_after = Pt(4 if indice == 0 else 10)

        _adicionar_runs_formatados(
            paragrafo,
            bloco,
            negrito_base=indice == 0,
            tamanho=14 if indice == 0 else 12,
        )
        if indice == 0:
            _adicionar_linha_dupla_docx(documento)
        primeira_linha = False

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
